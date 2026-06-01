#!/usr/bin/env python3
"""Baut für jede Testakte ein 'gesamt-pdf/<name>_gesamt.pdf', das alle
Aktenstücke (Markdown, TXT, EML, CSV, XLSX, DOCX, Bilder, PDF) in ein einziges,
sauber gerendertes Dokument mit Cover, Inhaltsverzeichnis und Seitenzahlen
zusammenfasst.

Aufruf:
  python3 scripts/build-testakte-gesamt-pdf.py                 # alle Testakten
  python3 scripts/build-testakte-gesamt-pdf.py <name1> <name2>  # gezielt
"""

from __future__ import annotations

import io
import re
import sys
import csv
from email import policy
from email.parser import BytesParser
from pathlib import Path

# Drittabhaengigkeiten
from openpyxl import load_workbook
from pypdf import PdfReader, PdfWriter
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import cm
from reportlab.lib.colors import HexColor, black
from reportlab.platypus import (
    SimpleDocTemplate,
    Image as RLImage,
    Paragraph,
    Spacer,
    PageBreak,
    Table,
    TableStyle,
)
from reportlab.lib.utils import ImageReader
from reportlab.lib.enums import TA_LEFT
from reportlab.pdfgen import canvas
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont

# DOCX
try:
    from docx import Document
except ImportError:
    Document = None  # type: ignore

REPO_ROOT = Path(__file__).resolve().parent.parent
TESTAKTEN = REPO_ROOT / "testakten"

# Design
TEAL = HexColor("#01696F")
MUTED = HexColor("#7A7974")
BORDER = HexColor("#D4D1CA")
SURFACE = HexColor("#F7F6F2")

# Font: System-Helvetica als Fallback. Inter waere schoener, aber wir verzichten
# auf Netzwerk-Downloads, damit das Skript offline laeuft.
FONT_REG = "Helvetica"
FONT_BOLD = "Helvetica-Bold"

styles = getSampleStyleSheet()
s_cover_label = ParagraphStyle(
    "CoverLabel",
    fontName=FONT_REG, fontSize=14, leading=18,
    textColor=MUTED, spaceAfter=6,
)
s_cover_title = ParagraphStyle(
    "CoverTitle",
    fontName=FONT_BOLD, fontSize=28, leading=34,
    textColor=TEAL, alignment=TA_LEFT, spaceAfter=14,
)
s_cover_sub = ParagraphStyle(
    "CoverSub",
    fontName=FONT_REG, fontSize=12, leading=16,
    textColor=black, spaceAfter=4,
)
s_cover_meta = ParagraphStyle(
    "CoverMeta",
    fontName=FONT_REG, fontSize=9, leading=12,
    textColor=MUTED, spaceAfter=3,
)
s_h1 = ParagraphStyle(
    "H1", parent=styles["Heading1"],
    fontName=FONT_BOLD, fontSize=18, leading=22, textColor=TEAL,
    spaceBefore=18, spaceAfter=8,
)
s_h2 = ParagraphStyle(
    "H2", parent=styles["Heading2"],
    fontName=FONT_BOLD, fontSize=14, leading=18, textColor=black,
    spaceBefore=12, spaceAfter=6,
)
s_h3 = ParagraphStyle(
    "H3", parent=styles["Heading3"],
    fontName=FONT_BOLD, fontSize=11, leading=14, textColor=black,
    spaceBefore=8, spaceAfter=4,
)
s_body = ParagraphStyle(
    "Body", parent=styles["BodyText"],
    fontName=FONT_REG, fontSize=10, leading=14, textColor=black,
    spaceAfter=6,
)
s_table_compact = ParagraphStyle(
    "TableCompact", parent=styles["BodyText"],
    fontName=FONT_REG, fontSize=7.5, leading=9, textColor=black,
    spaceAfter=0,
)
s_meta = ParagraphStyle(
    "Meta", parent=styles["BodyText"],
    fontName=FONT_REG, fontSize=9, leading=12, textColor=MUTED,
    spaceAfter=4,
)
s_partlabel = ParagraphStyle(
    "PartLabel", parent=styles["BodyText"],
    fontName=FONT_BOLD, fontSize=11, leading=14, textColor=MUTED,
    spaceAfter=2,
)

# Briefkopf-/Vermerks-Styles (authentischere Darstellung)
BANNER_BEHOERDE = HexColor("#0E3A5B")     # dunkles Behoerden-Blau (BaFin-naher Ton)
BANNER_KANZLEI = HexColor("#1F3F2E")      # tiefes Kanzlei-Gruen
BANNER_INTERN = HexColor("#8A1C1C")       # warnfarbenes Rot fuer interne Vermerke
BANNER_BANK = HexColor("#23344F")         # gedaempftes Bank-Blau
BANNER_ENTWURF = HexColor("#7A6A1E")      # Ocker fuer Entwuerfe / Presse
BANNER_INTERN_DOK = HexColor("#4B3B6B")   # gedaempftes Lila fuer interne Hauspapiere
BANNER_BILD = HexColor("#4A4A45")         # neutrales Grau fuer Bilddokumentation

s_letterhead_org = ParagraphStyle(
    "LetterheadOrg",
    fontName=FONT_BOLD, fontSize=13, leading=16, textColor=black,
    spaceAfter=2,
)
s_letterhead_unit = ParagraphStyle(
    "LetterheadUnit",
    fontName=FONT_REG, fontSize=9, leading=12, textColor=MUTED,
    spaceAfter=2,
)
s_letterhead_addr = ParagraphStyle(
    "LetterheadAddr",
    fontName=FONT_REG, fontSize=9, leading=12, textColor=black,
    spaceAfter=2,
)
s_banner_text = ParagraphStyle(
    "BannerText",
    fontName=FONT_BOLD, fontSize=9, leading=11, textColor=HexColor("#FFFFFF"),
    spaceAfter=0,
)
s_stamp = ParagraphStyle(
    "Stamp",
    fontName=FONT_BOLD, fontSize=9, leading=11, textColor=HexColor("#8A1C1C"),
    spaceAfter=2,
)
s_betreff = ParagraphStyle(
    "Betreff",
    fontName=FONT_BOLD, fontSize=11, leading=14, textColor=black,
    spaceBefore=6, spaceAfter=4,
)

# Reihenfolge der Datei-Typen im Gesamt-PDF
TYPE_ORDER = ["md", "txt", "eml", "csv", "xlsx", "docx", "image", "pdf"]
IMAGE_EXTS = {"jpg", "jpeg", "png"}
TYPE_LABEL = {
    "md": "Aktenstücke (Markdown)",
    "txt": "Notizen und Textdateien",
    "eml": "E-Mails",
    "csv": "CSV-Tabellen",
    "xlsx": "Excel-Tabellen",
    "docx": "Word-Dokumente",
    "image": "Bildanlagen und Screenshots",
    "pdf": "PDF-Anhänge (Originaldokumente)",
}


def escape(s: str) -> str:
    return s.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")


# ---------------------------------------------------------------------------
# AKTE-META: Briefkopf-/Stuecktyp-Block am Anfang einer Markdown-Datei
# ---------------------------------------------------------------------------
# Format (HTML-Kommentar, kein YAML-Frontmatter):
#
#   <!-- AKTE-META
#   typ: original_eingehend | original_ausgehend | interner_vermerk | mandantenkorrespondenz | entwurf
#   absender: Name; Einheit; Strasse; PLZ Ort
#   adressat: Name; Einheit; Strasse; PLZ Ort
#   datum: 18. Maerz 2026
#   az: VBS 4 1 7-K-22-188/2026
#   ihr_zeichen: SR-2026-FIN-0612
#   betreff: Anordnung Sonderpruefung gem. Paragraph 44 KWG
#   eingangsstempel: 18.03.2026 09:47 Uhr (Einschreiben Rueckschein)  # nur original_eingehend
#   vertraulichkeit: VERTRAULICH | ANWALTLICH PRIVILEGIERT | INTERN  # optional
#   -->
#
# Jede Zeile innerhalb des Kommentars hat 'schluessel: wert'.
# Mehrzeilige Werte sind erlaubt durch Einrueckung der Folgezeilen.

AKTE_META_RE = re.compile(r"<!--\s*AKTE-META\s*\n(.*?)\n\s*-->\s*\n?", re.DOTALL)

TYP_LABEL = {
    "original_eingehend": "ORIGINAL-SCHREIBEN (eingegangen)",
    "original_ausgehend": "ORIGINAL-SCHREIBEN (versandt)",
    "interner_vermerk": "INTERNER VERMERK",
    "mandantenkorrespondenz": "MANDANTENKORRESPONDENZ",
    "entwurf": "ENTWURF",
    "internes_dokument": "INTERNES HAUSDOKUMENT",
    "bild_dokumentation": "BILD-DOKUMENTATION",
}

TYP_BANNER = {
    "original_eingehend": BANNER_BEHOERDE,
    "original_ausgehend": BANNER_KANZLEI,
    "interner_vermerk": BANNER_INTERN,
    "mandantenkorrespondenz": BANNER_BANK,
    "entwurf": BANNER_ENTWURF,
    "internes_dokument": BANNER_INTERN_DOK,
    "bild_dokumentation": BANNER_BILD,
}


def parse_sidecar_meta(path: Path) -> dict:
    """Liest <datei>.meta neben einer Binaer-/EML-Beilage.

    Format identisch zum AKTE-META-Block, aber ohne HTML-Kommentar-Klammer:
      typ: original_eingehend
      absender: ...
      adressat: ...
      ...
    Gibt {} zurueck, wenn keine Sidecar existiert.
    """
    sidecar = path.with_suffix(path.suffix + ".meta")
    if not sidecar.exists():
        return {}
    meta: dict = {}
    current_key = None
    for raw in sidecar.read_text(encoding="utf-8", errors="replace").splitlines():
        if not raw.strip() or raw.lstrip().startswith("#"):
            continue
        if raw.startswith(" ") or raw.startswith("\t"):
            if current_key:
                meta[current_key] = (meta.get(current_key, "") + " " + raw.strip()).strip()
            continue
        if ":" in raw:
            k, _, v = raw.partition(":")
            current_key = k.strip().lower()
            meta[current_key] = v.strip()
    return meta


def parse_akte_meta(md_text: str) -> tuple[dict, str]:
    m = AKTE_META_RE.search(md_text)
    if not m:
        return {}, md_text
    block = m.group(1)
    meta: dict = {}
    current_key = None
    for raw in block.splitlines():
        if not raw.strip():
            continue
        if raw.startswith(" ") or raw.startswith("\t"):
            if current_key:
                meta[current_key] = (meta.get(current_key, "") + " " + raw.strip()).strip()
            continue
        if ":" in raw:
            k, _, v = raw.partition(":")
            current_key = k.strip().lower()
            meta[current_key] = v.strip()
    body = md_text[: m.start()] + md_text[m.end() :]
    return meta, body


def _addr_block(value: str, label: str) -> list:
    """Adressblock aus 'Name; Einheit; Strasse; PLZ Ort' -> Flowables."""
    out = [Paragraph(f"<b>{escape(label)}</b>", s_letterhead_unit)]
    parts = [p.strip() for p in value.split(";") if p.strip()]
    if not parts:
        return out
    out.append(Paragraph(escape(parts[0]), s_letterhead_org))
    for p in parts[1:]:
        out.append(Paragraph(escape(p), s_letterhead_addr))
    return out


def render_briefkopf(meta: dict) -> list:
    """Erzeugt eine getypte Briefkopf-/Vermerk-Box vor dem Stueck-Body."""
    typ = (meta.get("typ") or "").strip().lower()
    if typ not in TYP_LABEL:
        return []

    label = TYP_LABEL[typ]
    banner_color = TYP_BANNER[typ]
    vertraulichkeit = (meta.get("vertraulichkeit") or "").strip()

    out: list = []

    # Banner-Zeile (farbig, weisser Text) - signalisiert Stuecktyp sofort
    banner_text = label
    if vertraulichkeit:
        banner_text = label + "   |   " + vertraulichkeit
    banner_tbl = Table(
        [[Paragraph(escape(banner_text), s_banner_text)]],
        colWidths=[16 * cm],
    )
    banner_tbl.setStyle(TableStyle([
        ("BACKGROUND", (0, 0), (-1, -1), banner_color),
        ("LEFTPADDING", (0, 0), (-1, -1), 8),
        ("RIGHTPADDING", (0, 0), (-1, -1), 8),
        ("TOPPADDING", (0, 0), (-1, -1), 5),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 5),
    ]))
    out.append(banner_tbl)
    out.append(Spacer(1, 4))

    # Adressblock-Tabelle (links Absender, rechts Adressat)
    absender = meta.get("absender", "")
    adressat = meta.get("adressat", "")
    if absender or adressat:
        left_cell = _addr_block(absender, "Absender") if absender else [Paragraph("", s_letterhead_addr)]
        right_cell = _addr_block(adressat, "Adressat") if adressat else [Paragraph("", s_letterhead_addr)]
        addr_tbl = Table(
            [[left_cell, right_cell]],
            colWidths=[8 * cm, 8 * cm],
        )
        addr_tbl.setStyle(TableStyle([
            ("VALIGN", (0, 0), (-1, -1), "TOP"),
            ("LEFTPADDING", (0, 0), (-1, -1), 6),
            ("RIGHTPADDING", (0, 0), (-1, -1), 6),
            ("TOPPADDING", (0, 0), (-1, -1), 6),
            ("BOTTOMPADDING", (0, 0), (-1, -1), 6),
            ("BOX", (0, 0), (-1, -1), 0.4, BORDER),
            ("LINEAFTER", (0, 0), (0, -1), 0.3, BORDER),
            ("BACKGROUND", (0, 0), (-1, -1), HexColor("#FBFAF6")),
        ]))
        out.append(addr_tbl)
        out.append(Spacer(1, 4))

    # Metadaten-Tabelle: Datum, Az., Ihr Zeichen
    meta_rows = []
    if meta.get("datum"):
        meta_rows.append(["Datum", meta["datum"]])
    if meta.get("az"):
        meta_rows.append(["Aktenzeichen", meta["az"]])
    if meta.get("ihr_zeichen"):
        meta_rows.append(["Ihr Zeichen", meta["ihr_zeichen"]])
    if meta_rows:
        mtbl = Table(
            [[Paragraph(f"<b>{escape(k)}</b>", s_letterhead_addr),
              Paragraph(escape(v), s_letterhead_addr)] for k, v in meta_rows],
            colWidths=[3.5 * cm, 12.5 * cm],
        )
        mtbl.setStyle(TableStyle([
            ("VALIGN", (0, 0), (-1, -1), "TOP"),
            ("LEFTPADDING", (0, 0), (-1, -1), 6),
            ("RIGHTPADDING", (0, 0), (-1, -1), 6),
            ("TOPPADDING", (0, 0), (-1, -1), 2),
            ("BOTTOMPADDING", (0, 0), (-1, -1), 2),
        ]))
        out.append(mtbl)

    # Eingangsstempel nur bei eingegangenen Originalen
    if typ == "original_eingehend" and meta.get("eingangsstempel"):
        stamp_tbl = Table(
            [[Paragraph(
                "EINGEGANGEN<br/>" + escape(meta["eingangsstempel"]),
                s_stamp,
            )]],
            colWidths=[6 * cm],
            hAlign="RIGHT",
        )
        stamp_tbl.setStyle(TableStyle([
            ("BOX", (0, 0), (-1, -1), 1.0, HexColor("#8A1C1C")),
            ("LEFTPADDING", (0, 0), (-1, -1), 6),
            ("RIGHTPADDING", (0, 0), (-1, -1), 6),
            ("TOPPADDING", (0, 0), (-1, -1), 4),
            ("BOTTOMPADDING", (0, 0), (-1, -1), 4),
        ]))
        out.append(Spacer(1, 2))
        out.append(stamp_tbl)

    # Betreff
    if meta.get("betreff"):
        out.append(Paragraph("Betreff: " + escape(meta["betreff"]), s_betreff))

    out.append(Spacer(1, 8))
    return out


def md_to_flowables(md_text: str) -> list:
    out: list = []
    # Briefkopf-/Vermerk-Block aus AKTE-META vor dem eigentlichen Body
    meta, md_text = parse_akte_meta(md_text)
    if meta:
        out.extend(render_briefkopf(meta))
    lines = md_text.splitlines()
    i = 0
    while i < len(lines):
        line = lines[i].rstrip()
        if not line:
            i += 1
            continue
        if line.startswith("# "):
            out.append(Paragraph(escape(line[2:].strip()), s_h1))
            i += 1
            continue
        if line.startswith("## "):
            out.append(Paragraph(escape(line[3:].strip()), s_h2))
            i += 1
            continue
        if line.startswith("### "):
            out.append(Paragraph(escape(line[4:].strip()), s_h3))
            i += 1
            continue
        if line.startswith("---"):
            out.append(Spacer(1, 6))
            i += 1
            continue
        # Tabelle?
        if (
            line.startswith("|")
            and i + 1 < len(lines)
            and re.match(r"^\|[\s\-:|]+\|$", lines[i + 1])
        ):
            header = [c.strip() for c in line.strip("|").split("|")]
            i += 2
            rows = [header]
            while i < len(lines) and lines[i].startswith("|"):
                cells = [c.strip() for c in lines[i].strip("|").split("|")]
                rows.append(cells)
                i += 1
            col_count = max(1, len(header))
            avail_width = 16 * cm
            col_widths = [avail_width / col_count] * col_count
            if col_count >= 2:
                col_widths[0] = min(4 * cm, avail_width / col_count * 1.5)
                rest = (avail_width - col_widths[0]) / (col_count - 1)
                for k in range(1, col_count):
                    col_widths[k] = rest
            # Bei breiten Tabellen (6+ Spalten) kompakten Style nutzen, sonst sprengt der Zellinhalt die Seite
            cell_style = s_table_compact if col_count >= 6 else s_body
            tbl_data = []
            for r in rows:
                tbl_data.append([Paragraph(escape(c), cell_style) for c in r])
            tbl = Table(tbl_data, colWidths=col_widths, repeatRows=1, splitByRow=1)
            tbl.setStyle(
                TableStyle(
                    [
                        ("BACKGROUND", (0, 0), (-1, 0), SURFACE),
                        ("FONTNAME", (0, 0), (-1, 0), FONT_BOLD),
                        ("BOX", (0, 0), (-1, -1), 0.4, BORDER),
                        ("INNERGRID", (0, 0), (-1, -1), 0.3, BORDER),
                        ("VALIGN", (0, 0), (-1, -1), "TOP"),
                        ("LEFTPADDING", (0, 0), (-1, -1), 4),
                        ("RIGHTPADDING", (0, 0), (-1, -1), 4),
                        ("TOPPADDING", (0, 0), (-1, -1), 3),
                        ("BOTTOMPADDING", (0, 0), (-1, -1), 3),
                    ]
                )
            )
            out.append(tbl)
            out.append(Spacer(1, 6))
            continue
        if line.startswith("- ") or line.startswith("* "):
            text = line[2:].strip()
            text = re.sub(r"\*\*([^*]+)\*\*", r"<b>\1</b>", text)
            out.append(Paragraph("• " + _inline_markup(text), s_body))
            i += 1
            continue
        if re.match(r"^\d+\.\s", line):
            text = re.sub(r"\*\*([^*]+)\*\*", r"<b>\1</b>", line)
            out.append(Paragraph(_inline_markup(text), s_body))
            i += 1
            continue
        # Sammle normalen Absatz bis zur naechsten Leerzeile/Sondersyntax
        block = [line]
        j = i + 1
        while (
            j < len(lines)
            and lines[j].strip()
            and not lines[j].startswith(("#", "-", "*", "|", "---"))
            and not re.match(r"^\d+\.\s", lines[j])
        ):
            block.append(lines[j].rstrip())
            j += 1
        text = " ".join(block).strip()
        text = re.sub(r"\*\*([^*]+)\*\*", r"<b>\1</b>", text)
        text = re.sub(r"`([^`]+)`", r"<font face='Courier'>\1</font>", text)
        out.append(Paragraph(_inline_markup(text), s_body))
        i = j
    return out


def _inline_markup(s: str) -> str:
    """Escape minimal, lasse erlaubte Inline-Tags."""
    s = s.replace("&", "&amp;")
    # Erlaubte Tags wieder herstellen
    s = re.sub(r"&lt;(/?(?:b|i|sub|sup))&gt;", r"<\1>", s)
    s = s.replace("&lt;font face='Courier'&gt;", "<font face='Courier'>")
    s = s.replace("&lt;/font&gt;", "</font>")
    # Eckige Klammern in normalem Text: behalten, aber nicht als Tag
    return s


def txt_to_flowables(text: str) -> list:
    out = []
    for para in text.split("\n\n"):
        para = para.strip()
        if not para:
            continue
        # Zeilenumbrueche im Absatz erhalten
        out.append(Paragraph(escape(para).replace("\n", "<br/>"), s_body))
    return out


def eml_to_flowables(path: Path) -> list:
    out = []
    try:
        with open(path, "rb") as f:
            msg = BytesParser(policy=policy.default).parse(f)
        headers = [
            ("Von", msg.get("From", "")),
            ("An", msg.get("To", "")),
            ("Datum", msg.get("Date", "")),
            ("Betreff", msg.get("Subject", "")),
        ]
        body_part = msg.get_body(preferencelist=("plain", "html"))
        body = body_part.get_content() if body_part else ""
    except Exception as e:
        out.append(Paragraph(f"<i>E-Mail konnte nicht gelesen werden: {escape(str(e))}</i>", s_meta))
        return out

    rows = [
        [Paragraph(label, s_meta), Paragraph(escape(value), s_meta)]
        for label, value in headers
    ]
    tbl = Table(rows, colWidths=[2.5 * cm, 13.5 * cm])
    tbl.setStyle(
        TableStyle(
            [
                ("BACKGROUND", (0, 0), (0, -1), SURFACE),
                ("BOX", (0, 0), (-1, -1), 0.3, BORDER),
                ("INNERGRID", (0, 0), (-1, -1), 0.2, BORDER),
                ("VALIGN", (0, 0), (-1, -1), "TOP"),
                ("LEFTPADDING", (0, 0), (-1, -1), 4),
                ("RIGHTPADDING", (0, 0), (-1, -1), 4),
                ("TOPPADDING", (0, 0), (-1, -1), 3),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 3),
            ]
        )
    )
    out.append(tbl)
    out.append(Spacer(1, 6))
    out.extend(txt_to_flowables(body))
    return out


def csv_to_flowables(path: Path) -> list:
    out = []
    try:
        with open(path, encoding="utf-8") as f:
            reader = csv.reader(f)
            rows = list(reader)
    except UnicodeDecodeError:
        with open(path, encoding="latin-1") as f:
            reader = csv.reader(f)
            rows = list(reader)
    except Exception as e:
        out.append(Paragraph(f"<i>CSV konnte nicht gelesen werden: {escape(str(e))}</i>", s_meta))
        return out

    if not rows:
        return out
    max_cols = max(len(r) for r in rows)
    rows = [r + [""] * (max_cols - len(r)) for r in rows]
    out.extend(_render_table(rows, header=True))
    return out


def xlsx_to_flowables(path: Path) -> list:
    out = []
    try:
        wb = load_workbook(path, data_only=True)
    except Exception as e:
        out.append(Paragraph(f"<i>XLSX konnte nicht gelesen werden: {escape(str(e))}</i>", s_meta))
        return out
    for sheet_name in wb.sheetnames:
        ws = wb[sheet_name]
        out.append(Paragraph(f"Tabellenblatt: {escape(sheet_name)}", s_h3))
        rows = []
        for row in ws.iter_rows(values_only=True):
            rows.append(
                [_format_cell(c) for c in row]
            )
        if not rows:
            continue
        # Leere Zeilen/Spalten am Ende abschneiden
        while rows and not any(c.strip() for c in rows[-1]):
            rows.pop()
        if not rows:
            continue
        max_cols = max(len(r) for r in rows)
        if max_cols == 0:
            continue
        # Hinten leere Spalten abschneiden
        while max_cols > 0 and all(
            (len(r) <= max_cols - 1) or (not r[max_cols - 1].strip()) for r in rows
        ):
            max_cols -= 1
        if max_cols == 0:
            continue
        rows = [r[:max_cols] + [""] * (max_cols - len(r[:max_cols])) for r in rows]
        out.extend(_render_table(rows, header=True))
        out.append(Spacer(1, 6))
    return out


def _format_cell(c) -> str:
    if c is None:
        return ""
    if isinstance(c, float):
        if c == int(c):
            return str(int(c))
        return f"{c:.4f}".rstrip("0").rstrip(".")
    return str(c)


# Maximalzeichen pro Zelle, ab denen die Tabelle nicht mehr als Table gerendert wird,
# sondern als sequentielle Absatzfolge (verhindert ReportLab-Overflow).
_MAX_CELL_CHARS = 1200


def _split_long_text(text: str, chunk: int = 800) -> list:
    """Schneidet sehr langen Text an Absatz- oder Satzgrenzen in Stuecke."""
    text = text.replace("\r", "")
    if len(text) <= chunk:
        return [text]
    # Erst Absaetze probieren
    paras = [p for p in text.split("\n") if p.strip()]
    if any(len(p) > chunk for p in paras):
        # Weiter an Saetzen schneiden
        out = []
        for p in paras:
            if len(p) <= chunk:
                out.append(p)
                continue
            buf = ""
            for sent in p.replace("; ", "; |").replace(". ", ". |").split("|"):
                if len(buf) + len(sent) > chunk and buf:
                    out.append(buf)
                    buf = sent
                else:
                    buf += sent
            if buf:
                out.append(buf)
        return out
    return paras


def _render_table(rows: list, header: bool = False) -> list:
    """Rendert eine Tabelle. Falls Zellen zu lang werden, faellt es auf eine
    sequentielle Absatzdarstellung zurueck (Reihe fuer Reihe), damit ReportLab
    keine Overflow-Fehler wirft."""
    max_cell_len = max((len(c) for r in rows for c in r), default=0)
    max_cols_in_table = max((len(r) for r in rows), default=0)
    # Bei sehr breiten Tabellen (>12 Spalten) faellt es sequentiell zurueck,
    # weil die Spaltenbreite sonst kleiner ist als die kleinste Wortbreite
    # und ReportLab Cell-Overflow-Fehler wirft.
    if max_cell_len > _MAX_CELL_CHARS or max_cols_in_table > 12:
        out = []
        header_row = rows[0] if header else None
        body_rows = rows[1:] if header else rows
        for ri, r in enumerate(body_rows):
            if header_row:
                # Reihen-Trennlinie + Spaltenkopf pro Zelle
                for ci, cell in enumerate(r):
                    label = header_row[ci] if ci < len(header_row) else f"Spalte {ci+1}"
                    if label.strip():
                        out.append(Paragraph(f"<b>{escape(label)}</b>", s_meta))
                    for chunk in _split_long_text(cell):
                        out.append(Paragraph(escape(chunk), s_body))
            else:
                for ci, cell in enumerate(r):
                    for chunk in _split_long_text(cell):
                        out.append(Paragraph(escape(chunk), s_body))
            out.append(Spacer(1, 4))
            out.append(HRFlowable(width="100%", thickness=0.2, color=BORDER))
            out.append(Spacer(1, 4))
        return out

    max_cols = max(len(r) for r in rows)
    avail_width = 16 * cm
    col_widths = [avail_width / max_cols] * max_cols
    data = [
        [Paragraph(escape(c), s_meta) for c in r]
        for r in rows
    ]
    tbl = Table(data, colWidths=col_widths, repeatRows=1 if header else 0, splitByRow=1)
    cmds = [
        ("BOX", (0, 0), (-1, -1), 0.3, BORDER),
        ("INNERGRID", (0, 0), (-1, -1), 0.2, BORDER),
        ("VALIGN", (0, 0), (-1, -1), "TOP"),
        ("LEFTPADDING", (0, 0), (-1, -1), 3),
        ("RIGHTPADDING", (0, 0), (-1, -1), 3),
        ("TOPPADDING", (0, 0), (-1, -1), 2),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 2),
    ]
    if header:
        cmds.insert(0, ("BACKGROUND", (0, 0), (-1, 0), SURFACE))
        cmds.insert(1, ("FONTNAME", (0, 0), (-1, 0), FONT_BOLD))
    tbl.setStyle(TableStyle(cmds))
    return [tbl]


def docx_to_flowables(path: Path) -> list:
    out = []
    if Document is None:
        out.append(Paragraph("<i>python-docx nicht installiert, Inhalt wird uebersprungen.</i>", s_meta))
        return out
    try:
        doc = Document(str(path))
    except Exception as e:
        out.append(Paragraph(f"<i>DOCX konnte nicht gelesen werden: {escape(str(e))}</i>", s_meta))
        return out
    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue
        style = para.style.name if para.style else ""
        if style.startswith("Heading 1"):
            out.append(Paragraph(escape(text), s_h2))
        elif style.startswith("Heading 2"):
            out.append(Paragraph(escape(text), s_h3))
        elif style.startswith("Heading"):
            out.append(Paragraph(escape(text), s_h3))
        else:
            out.append(Paragraph(escape(text), s_body))
    for table in doc.tables:
        rows = []
        for r in table.rows:
            rows.append([c.text.strip() for c in r.cells])
        if rows:
            out.extend(_render_table(rows, header=True))
            out.append(Spacer(1, 6))
    return out


def image_to_flowables(path: Path) -> list:
    out = []
    try:
        width, height = ImageReader(str(path)).getSize()
        max_width = 16 * cm
        max_height = 22 * cm
        scale = min(max_width / width, max_height / height, 1)
        img = RLImage(str(path), width=width * scale, height=height * scale)
        out.append(img)
        out.append(Spacer(1, 4))
        out.append(Paragraph(f"Bilddatei: {escape(path.name)}", s_meta))
    except Exception as e:
        out.append(Paragraph(f"<i>Bild konnte nicht gerendert werden: {escape(str(e))}</i>", s_meta))
    return out


def header_footer_factory(testakte_name: str):
    def hf(canv: canvas.Canvas, doc) -> None:
        canv.saveState()
        canv.setFont(FONT_REG, 8)
        canv.setFillColor(MUTED)
        canv.drawString(2 * cm, 1.2 * cm, f"Arbeitsakte: {testakte_name}")
        canv.drawRightString(19 * cm, 1.2 * cm, f"Seite {doc.page}")
        canv.setStrokeColor(BORDER)
        canv.setLineWidth(0.3)
        canv.line(2 * cm, 1.6 * cm, 19 * cm, 1.6 * cm)
        canv.restoreState()

    return hf


def build_cover(name: str, readme_summary: str | None, h1: str | None = None) -> list:
    title = h1 if h1 else name
    out = [
        Spacer(1, 4 * cm),
        Paragraph("Arbeitsakte", s_cover_label),
        Paragraph(escape(title), s_cover_title),
        Paragraph(escape(name), s_cover_meta),
        Spacer(1, 0.6 * cm),
    ]
    if readme_summary:
        out.append(Paragraph(escape(readme_summary), s_cover_sub))
        out.append(Spacer(1, 0.8 * cm))
    else:
        out.append(Spacer(1, 0.8 * cm))
    out.append(Paragraph(
        "Diese Datei bündelt alle Aktenstücke in einem Dokument. "
        "Die Einzeldateien liegen im Aktenordner ebenfalls vor.",
        s_cover_meta,
    ))
    return out


def extract_readme_summary(readme_path: Path) -> tuple[str | None, str | None]:
    """Liest aus der README den H1-Titel und einen kurzen beschreibenden Absatz.

    Der beschreibende Absatz wird absichtlich erst gesucht, NACHDEM Download-Bloecke
    und Aktenstruktur-Blocks uebersprungen wurden, damit nicht der ZIP-Hinweis
    auf dem Cover landet.
    """
    if not readme_path.is_file():
        return None, None
    text = readme_path.read_text(encoding="utf-8")
    # H1
    h1_match = re.search(r"^#\s+(.+?)\s*$", text, flags=re.MULTILINE)
    h1 = h1_match.group(1).strip() if h1_match else None

    # Suche eine Sektion mit beschreibendem Inhalt: 'Kurzbild', 'Worum',
    # 'Sachverhalt', 'Ueberblick', 'Mandat', 'Fall' o.ae.
    section_pattern = re.compile(
        r"^##[^\n]*?(?:kurzbild|worum geht|sachverhalt|ueberblick|\u00fcberblick|mandat|fall|der fall|akte|kontext|ausgangslage|ausgangs|zweck|szenario|idee|einsatz|\u00fcbersicht|uebersicht|verfahrenseckdaten|aktenkern|aktenbestand|mandantenkonstellation|politische vorgabe|enthaltene arbeitsdateien|dateien)[^\n]*\n([\s\S]*?)(?=^## |\Z)",
        re.IGNORECASE | re.MULTILINE,
    )
    m = section_pattern.search(text)
    candidate_text = m.group(1) if m else text
    for para in candidate_text.split("\n\n"):
        para = para.strip()
        if not para:
            continue
        if para.startswith(("#", "-", "*", "|", "<!--", "```")):
            continue
        # Download-/ZIP-Hinweise ueberspringen
        lower = para.lower()
        if any(
            kw in lower
            for kw in (
                "zip-datei",
                "zip datei",
                "direkt-download",
                "als zip",
                "github-release",
                "github release",
                "download",
            )
        ):
            continue
        para = re.sub(r"\*\*([^*]+)\*\*", r"\1", para)
        para = re.sub(r"\s+", " ", para)
        return h1, para[:400]
    return h1, None


def collect_files(testakte_dir: Path) -> dict[str, list[Path]]:
    files_by_type: dict[str, list[Path]] = {t: [] for t in TYPE_ORDER}
    for f in testakte_dir.rglob("*"):
        if not f.is_file():
            continue
        # README und Gesamt-PDF ausschliessen
        if f.name == "README.md" and f.parent == testakte_dir:
            continue
        if "gesamt-pdf" in f.parts:
            continue
        ext = f.suffix.lower().lstrip(".")
        if ext in IMAGE_EXTS:
            files_by_type["image"].append(f)
            continue
        if ext not in TYPE_ORDER:
            continue
        files_by_type[ext].append(f)
    for t in files_by_type:
        files_by_type[t].sort(key=lambda p: str(p.relative_to(testakte_dir)).lower())
    return files_by_type


def build_text_pdf(testakte_dir: Path, files: dict[str, list[Path]], cover: list, tmp_path: Path) -> tuple[bool, list[Path]]:
    """Baut den Text-Teil als PDF, sammelt PDF-Anhaenge separat."""
    doc = SimpleDocTemplate(
        str(tmp_path),
        pagesize=A4,
        leftMargin=2 * cm, rightMargin=2 * cm,
        topMargin=2 * cm, bottomMargin=2 * cm,
        title=f"Arbeitsakte {testakte_dir.name}",
        author="Kanzleiakte",
    )
    flow = list(cover)
    flow.append(PageBreak())

    # Inhaltsverzeichnis (rudimentaer)
    toc_rows: list[list] = [["Teil", "Inhalt"]]
    teil_no = 1
    for t in TYPE_ORDER:
        if not files[t]:
            continue
        toc_rows.append([f"Teil {teil_no}", f"{TYPE_LABEL[t]} ({len(files[t])})"])
        teil_no += 1
    if len(toc_rows) > 1:
        flow.append(Paragraph("Inhaltsverzeichnis", s_h1))
        flow.append(Spacer(1, 8))
        flow.extend(_render_table(toc_rows, header=True))
        flow.append(PageBreak())

    pdf_attachments: list[Path] = []
    teil_no = 1
    for t in TYPE_ORDER:
        if not files[t]:
            continue
        if t == "pdf":
            # PDFs werden separat angehaengt (Original-Layout bewahren)
            pdf_attachments = files[t]
            teil_no += 1
            continue
        flow.append(Paragraph(f"Teil {teil_no} — {TYPE_LABEL[t]}", s_partlabel))
        flow.append(Paragraph(TYPE_LABEL[t], s_h1))
        flow.append(Spacer(1, 4))
        for f in files[t]:
            rel = f.relative_to(testakte_dir)
            flow.append(Paragraph(f"<b>Datei:</b> {escape(str(rel))}", s_meta))
            flow.append(Spacer(1, 4))
            # Sidecar-Briefkopf vor Nicht-MD-Beilagen (DOCX/EML/XLSX/CSV/TXT/JPG)
            # MD-Stuecke haben ihren AKTE-META-Block im Body und brauchen das nicht.
            if t != "md":
                sidecar = parse_sidecar_meta(f)
                if sidecar:
                    flow.extend(render_briefkopf(sidecar))
            try:
                if t == "md":
                    flow.extend(md_to_flowables(f.read_text(encoding="utf-8", errors="replace")))
                elif t == "txt":
                    flow.extend(txt_to_flowables(f.read_text(encoding="utf-8", errors="replace")))
                elif t == "eml":
                    flow.extend(eml_to_flowables(f))
                elif t == "csv":
                    flow.extend(csv_to_flowables(f))
                elif t == "xlsx":
                    flow.extend(xlsx_to_flowables(f))
                elif t == "docx":
                    flow.extend(docx_to_flowables(f))
                elif t == "image":
                    flow.extend(image_to_flowables(f))
            except Exception as e:
                flow.append(Paragraph(f"<i>Inhalt konnte nicht gerendert werden: {escape(str(e))}</i>", s_meta))
            flow.append(Spacer(1, 14))
        flow.append(PageBreak())
        teil_no += 1

    if len(flow) == len(cover) + 1:
        # Nichts ausser Cover -> trotzdem bauen, aber Hinweis
        flow.append(Paragraph(
            "Diese Arbeitsakte enthält keine renderbaren Inhalte ausserhalb der angefuegten PDFs.",
            s_body,
        ))

    hf = header_footer_factory(testakte_dir.name)
    try:
        doc.build(flow, onFirstPage=hf, onLaterPages=hf)
    except Exception as e:
        print(f"  FEHLER beim Bauen: {e}")
        return False, pdf_attachments
    return True, pdf_attachments


def append_pdf_with_separator(writer: PdfWriter, label: str, pdf_path: Path, testakte_name: str) -> None:
    sep = io.BytesIO()
    c = canvas.Canvas(sep, pagesize=A4)
    c.setTitle(label)
    c.setAuthor("Kanzleiakte")
    c.setFont(FONT_BOLD, 16)
    c.setFillColor(TEAL)
    c.drawString(2 * cm, 25 * cm, label)
    c.setFont(FONT_REG, 9)
    c.setFillColor(MUTED)
    c.drawString(2 * cm, 24.2 * cm, f"Datei: {pdf_path.name}")

    # Sidecar-Briefkopf auch fuer PDF-Anhaenge (auf der Separator-Seite)
    meta = parse_sidecar_meta(pdf_path)
    typ = (meta.get("typ") or "").strip().lower() if meta else ""
    if typ in TYP_LABEL:
        banner_color = TYP_BANNER[typ]
        banner_label = TYP_LABEL[typ]
        vertraulichkeit = (meta.get("vertraulichkeit") or "").strip()
        text = banner_label + ("   |   " + vertraulichkeit if vertraulichkeit else "")
        # farbiger Banner
        c.setFillColor(banner_color)
        c.rect(2 * cm, 22.5 * cm, 17 * cm, 0.9 * cm, fill=1, stroke=0)
        c.setFillColor(HexColor("#FFFFFF"))
        c.setFont(FONT_BOLD, 10)
        c.drawString(2.3 * cm, 22.85 * cm, text)
        # Absender / Adressat / Datum / Az.
        c.setFillColor(black)
        c.setFont(FONT_BOLD, 9)
        y = 21.6 * cm
        # Wert-Spalte bei 4.5 cm, damit lange Labels wie "Aktenzeichen:"
        # nicht in die Wert-Spalte hineinlaufen (9pt Helvetica-Bold ca. 2.15 cm)
        VALUE_X = 4.5 * cm
        if meta.get("absender"):
            c.drawString(2 * cm, y, "Absender:")
            c.setFont(FONT_REG, 9)
            for line in [p.strip() for p in meta["absender"].split(";")][:4]:
                c.drawString(VALUE_X, y, line)
                y -= 0.45 * cm
            c.setFont(FONT_BOLD, 9)
        if meta.get("adressat"):
            y -= 0.3 * cm
            c.drawString(2 * cm, y, "Adressat:")
            c.setFont(FONT_REG, 9)
            for line in [p.strip() for p in meta["adressat"].split(";")][:4]:
                c.drawString(VALUE_X, y, line)
                y -= 0.45 * cm
            c.setFont(FONT_BOLD, 9)
        y -= 0.3 * cm
        for key, lbl in [("datum", "Datum:"), ("az", "Aktenzeichen:"), ("ihr_zeichen", "Ihr Zeichen:")]:
            if meta.get(key):
                c.drawString(2 * cm, y, lbl)
                c.setFont(FONT_REG, 9)
                c.drawString(VALUE_X, y, meta[key])
                c.setFont(FONT_BOLD, 9)
                y -= 0.45 * cm
        if meta.get("betreff"):
            y -= 0.2 * cm
            c.drawString(2 * cm, y, "Betreff:")
            c.setFont(FONT_REG, 9)
            # Betreff ggf. umbrechen
            betreff = meta["betreff"]
            words = betreff.split()
            line = ""
            xb = VALUE_X
            for w in words:
                if len(line + " " + w) > 80:
                    c.drawString(xb, y, line.strip())
                    y -= 0.45 * cm
                    line = w
                else:
                    line = (line + " " + w).strip()
            if line:
                c.drawString(xb, y, line)
                y -= 0.45 * cm
        # Eingangsstempel-Box, falls eingehend
        if typ == "original_eingehend" and meta.get("eingangsstempel"):
            # Wort-genaues Wrapping per stringWidth statt Index-Slicing
            stamp_text = meta["eingangsstempel"]
            font_name = FONT_REG
            font_size = 8
            max_width = 5.6 * cm  # Box 6 cm minus 2x 0.2 cm Padding
            words = stamp_text.split()
            lines: list = []
            cur = ""
            for w in words:
                candidate = (cur + " " + w).strip()
                if pdfmetrics.stringWidth(candidate, font_name, font_size) <= max_width:
                    cur = candidate
                else:
                    if cur:
                        lines.append(cur)
                    # Wort selbst zu lang? Hartes Bruchstueck als Fallback
                    if pdfmetrics.stringWidth(w, font_name, font_size) > max_width:
                        chunk = ""
                        for ch in w:
                            if pdfmetrics.stringWidth(chunk + ch, font_name, font_size) <= max_width:
                                chunk += ch
                            else:
                                lines.append(chunk)
                                chunk = ch
                        cur = chunk
                    else:
                        cur = w
            if cur:
                lines.append(cur)
            # Box-Hoehe dynamisch: 1 Header-Zeile (0.55 cm) + n Stempel-Zeilen (0.4 cm)
            line_h = 0.4 * cm
            box_h = 0.7 * cm + len(lines) * line_h + 0.2 * cm
            box_y = 21.3 * cm - box_h  # Oberkante bei 21.3 cm fix
            c.setStrokeColor(HexColor("#8A1C1C"))
            c.setLineWidth(1.0)
            c.rect(13 * cm, box_y, 6 * cm, box_h, fill=0, stroke=1)
            c.setFillColor(HexColor("#8A1C1C"))
            c.setFont(FONT_BOLD, 9)
            c.drawString(13.2 * cm, box_y + box_h - 0.55 * cm, "EINGEGANGEN")
            c.setFont(font_name, font_size)
            ty = box_y + box_h - 0.55 * cm - 0.5 * cm
            for ln in lines:
                c.drawString(13.2 * cm, ty, ln)
                ty -= line_h

    c.setStrokeColor(BORDER)
    c.setLineWidth(0.3)
    c.line(2 * cm, 1.6 * cm, 19 * cm, 1.6 * cm)
    c.setFont(FONT_REG, 8)
    c.setFillColor(MUTED)
    c.drawString(2 * cm, 1.2 * cm, f"Arbeitsakte: {testakte_name}")
    c.showPage()
    c.save()
    sep.seek(0)
    for p in PdfReader(sep).pages:
        writer.add_page(p)
    try:
        for p in PdfReader(str(pdf_path)).pages:
            writer.add_page(p)
    except Exception as e:
        # PDF defekt oder verschluesselt -> Hinweisseite einfuegen
        sep2 = io.BytesIO()
        c2 = canvas.Canvas(sep2, pagesize=A4)
        c2.setFont(FONT_REG, 10)
        c2.drawString(2 * cm, 25 * cm, f"PDF konnte nicht eingebunden werden: {e}")
        c2.showPage()
        c2.save()
        sep2.seek(0)
        for p in PdfReader(sep2).pages:
            writer.add_page(p)


def build_gesamt_pdf(testakte_dir: Path) -> tuple[str, str]:
    """Gibt (status, info) zurueck. status in {ok, skip, error}."""
    name = testakte_dir.name
    out_dir = testakte_dir / "gesamt-pdf"
    out_dir.mkdir(exist_ok=True)
    out_path = out_dir / f"{name}_gesamt.pdf"

    files = collect_files(testakte_dir)
    total_files = sum(len(v) for v in files.values())
    if total_files == 0:
        return "skip", "keine Quelldateien"

    h1, summary = extract_readme_summary(testakte_dir / "README.md")
    cover = build_cover(name, summary, h1)

    tmp_text = Path(f"/tmp/_gesamt_text_{name}.pdf")
    ok, pdf_attachments = build_text_pdf(testakte_dir, files, cover, tmp_text)
    if not ok:
        return "error", "Text-PDF konnte nicht erzeugt werden"

    writer = PdfWriter()
    try:
        for page in PdfReader(str(tmp_text)).pages:
            writer.add_page(page)
    except Exception as e:
        return "error", f"Text-PDF nicht lesbar: {e}"

    for pdf in pdf_attachments:
        rel = pdf.relative_to(testakte_dir)
        label = f"PDF-Anhang: {rel}"
        append_pdf_with_separator(writer, label, pdf, name)

    writer.add_metadata(
        {
            "/Title": f"Arbeitsakte {name}",
            "/Author": "Kanzleiakte",
            "/Subject": "Gesamtakte",
        }
    )
    with open(out_path, "wb") as f:
        writer.write(f)
    size_kb = out_path.stat().st_size / 1024
    try:
        tmp_text.unlink()
    except Exception:
        pass
    return "ok", f"{out_path.relative_to(REPO_ROOT)} ({size_kb:.0f} KB, {total_files} Quelldateien)"


def main() -> None:
    targets = sys.argv[1:]
    all_dirs = sorted([d for d in TESTAKTEN.iterdir() if d.is_dir()])
    if targets:
        all_dirs = [d for d in all_dirs if d.name in targets]
    print(f"Verarbeite {len(all_dirs)} Testakten")
    print()
    counts = {"ok": 0, "skip": 0, "error": 0}
    for d in all_dirs:
        status, info = build_gesamt_pdf(d)
        counts[status] += 1
        sigil = {"ok": "OK ", "skip": "SK ", "error": "ERR"}[status]
        print(f"  {sigil} {d.name}: {info}")
    print()
    print(f"Fertig: {counts['ok']} OK, {counts['skip']} skip, {counts['error']} Fehler")


if __name__ == "__main__":
    main()
